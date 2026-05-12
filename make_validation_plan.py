#!/usr/bin/env python3
"""Run this script to generate the Word validation plan document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Biological_Validation_Plan.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("⚑  " + text)
    run.font.size    = Pt(10)
    run.font.italic  = True
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_table(headers, rows_data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row_data in rows_data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            row[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Biological Validation Plan")
run.font.size  = Pt(26)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run(
    "Systematic Precision / Recall / F1 Evaluation of the Forward "
    "Boolean-Network Regulatory Analysis Program"
)
run2.font.size   = Pt(14)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Arabidopsis thaliana Gene Regulatory Network  |  Bachelor Project 2026"
             ).font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1.  Overview and Objective")

body(
    "This plan describes how to systematically evaluate the biological accuracy of the "
    "forward Boolean network (BN) analysis program developed in this bachelor project. "
    "The program predicts which genes are activated or suppressed downstream of a given "
    "source transcription factor, using a knowledge-graph-derived Boolean network model "
    "with synchronous update semantics and a four-condition (dark/permissive × "
    "resting/perturbed) attractor comparison."
)
body(
    "The validation answers three questions that a peer reviewer will ask:"
)
bullet("Is the overlap between program predictions and experimentally known targets "
       "statistically greater than random?")
bullet("At which hop depth does the best trade-off between precision and recall occur?")
bullet("Do context-independent (robust) predictions outperform condition-specific "
       "(permissive-only) predictions in precision?")

body(
    "Two dedicated Python scripts have been written to automate the entire pipeline:"
)
bullet("eval_fw_predictions.py  —  runs the BN analysis for 19 test genes at hops 1–5 "
       "and saves all prediction sets to a single JSON file.")
bullet("eval_fw_metrics.py  —  loads the JSON, builds the ground truth, computes all "
       "metrics and FDR-corrected p-values, and writes two CSV tables.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — TEST GENE SELECTION
# ══════════════════════════════════════════════════════════════════════════════
h1("2.  Test Gene Selection")

body(
    "Test genes were selected by intersecting the raw ground-truth file "
    "(Ground_truth_experiment.csv, 207 930 rows) with the knowledge-graph network "
    "(filtered_networkL_normalized.csv). A gene qualifies as a test gene if:"
)
bullet("It is present as a source gene in the network.")
bullet("At least 3 of its experimentally confirmed downstream targets are also present "
       "in the network (ensuring enough ground-truth positives to compute meaningful metrics).")

body("The 19 selected test genes are:")

add_table(
    ["Gene", "Known activation targets in network", "Known suppression targets"],
    [
        ["pif4",   "55", "—"],
        ["EIN3",   "67", "3"],
        ["hy5",    "70", "5"],
        ["cop1",   "35", "10"],
        ["WRKY33", "37", "—"],
        ["GRF1",   "62", "—"],
        ["GRF3",   "56", "—"],
        ["MYB46",  "10", "—"],
        ["ABI4",   "37", "6"],
        ["SOG1",   "34", "3"],
        ["CCA1",   "~8", "~6"],
        ["ARR1",   "~10", "1"],
        ["ARF7",   "~9",  "1"],
        ["SOC1",   "~9",  "5"],
        ["SPL9",   "~9",  "1"],
        ["myc2",   "54",  "—"],
        ["ref6",   "33",  "—"],
        ["Abi5",   "37",  "—"],
        ["bzr1",   "28",  "—"],
    ],
    col_widths=[3.5, 5.5, 5.5]
)

note(
    "MYB46 is the primary gene of biological interest in this project. "
    "Its 10 known targets (including CSLA9 and secondary cell wall biosynthesis genes) "
    "provide a focused validation case even with a smaller ground-truth set. "
    "The remaining 18 genes expand the evaluation to diverse regulatory contexts "
    "and improve the statistical power of aggregate metrics."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — GROUND TRUTH CONSTRUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1("3.  Ground Truth Construction")

body(
    "The ground truth is extracted from Ground_truth_experiment.csv using keyword "
    "matching on the free-text 'relationship' column, then filtered to gene–gene "
    "pairs where both genes are present in the knowledge-graph network. "
    "The same clean_name() function used by the BN programs is applied to both sides "
    "to ensure name consistency."
)

h2("3.1  Activation ground truth keywords")
body("Rows matching any of the following (case-insensitive) are classified as "
     "activation relationships:")
code("activates | induces | upregulates | promotes | directly regulates")
code("regulates expression of | targets | increases expression")
code("binds.*promoter | is a target.*of | is regulated.*by")

h2("3.2  Suppression ground truth keywords")
code("represses | suppresses | inhibits | downregulates")
code("negatively regulates | reduces expression | is repressed by")

h2("3.3  Name normalisation")
body(
    "Gene names in the ground truth use inconsistent capitalisation (e.g. 'MYB46', "
    "'myb46', 'Myb46'). A case-insensitive lookup dictionary is built from the "
    "network's canonical names so that all three forms resolve to the same key. "
    "This step is critical: without it, true positives are silently counted as false "
    "positives and precision is artificially deflated."
)

h2("3.4  Incompleteness caveat")
body(
    "The literature-derived ground truth is necessarily incomplete — only relationships "
    "that have been experimentally studied and reported in the papers covered by the "
    "knowledge graph are included. This has a direct consequence for interpretation:"
)
bullet("Precision is an accurate metric: a predicted gene either has published "
       "experimental evidence or it does not.")
bullet("Recall is a lower bound: a predicted gene absent from the ground truth may "
       "still be a true regulatory target that has not yet been studied.")
body(
    "This asymmetry must be stated explicitly in the Methods section of the paper."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PREDICTION SETS
# ══════════════════════════════════════════════════════════════════════════════
h1("4.  Prediction Sets and Their Biological Meaning")

body(
    "At each (gene, hop) combination the program produces four activation and two "
    "suppression prediction sets, in descending order of expected precision:"
)

add_table(
    ["Prediction set", "Definition", "Expected property"],
    [
        ["robust_activated",
         "Stably OFF→ON in both dark and permissive perturbed attractors",
         "Highest precision — activation is context-independent"],
        ["perm_activated",
         "Stably OFF→ON in permissive perturbed attractors",
         "Main evaluation set; wider coverage than robust"],
        ["dark_activated",
         "Stably OFF→ON in dark perturbed attractors only",
         "More conservative starting condition"],
        ["conditional_derepressed",
         "OFF in some resting attractors, ON in some perturbed attractors",
         "Weakest signal; partial or context-dependent effect"],
        ["robust_suppressed",
         "Stably ON→OFF in both backgrounds",
         "Highest-confidence suppression"],
        ["perm_suppressed",
         "Stably ON→OFF in permissive background",
         "Main suppression evaluation set"],
    ],
    col_widths=[4.0, 6.5, 5.0]
)

body(
    "All six sets are stored in predictions.json and evaluated independently by "
    "eval_fw_metrics.py. The primary set for reporting in the paper is "
    "perm_activated, with robust_activated used as a precision-vs-recall comparison."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — METRICS
# ══════════════════════════════════════════════════════════════════════════════
h1("5.  Metrics and Statistical Framework")

h2("5.1  Precision, Recall, and F1")

body("For a given source gene G, prediction set P, and ground truth set GT:")

add_table(
    ["Quantity", "Formula", "Interpretation"],
    [
        ["True Positives (TP)", "|P ∩ GT|",
         "Predicted genes confirmed by literature"],
        ["False Positives (FP)", "|P \\ GT|",
         "Predicted genes with no published evidence"],
        ["False Negatives (FN)", "|GT \\ P|",
         "Known targets the program missed"],
        ["Precision", "TP / (TP + FP)",
         "Quality: fraction of predictions that are real"],
        ["Recall", "TP / (TP + FN)",
         "Coverage: fraction of known targets predicted (lower bound)"],
        ["F1", "2·P·R / (P+R)",
         "Harmonic mean — single balanced score"],
    ],
    col_widths=[3.8, 3.8, 7.9]
)

h2("5.2  Hypergeometric significance test")

body(
    "For each (gene, hop, prediction type) triple, a hypergeometric p-value tests "
    "whether the observed overlap between predictions and ground truth exceeds what "
    "would be expected by random sampling:"
)
bullet("Population N = number of network genes in the downstream subgraph at this hop")
bullet("Success states K = |GT| (ground-truth targets within the subgraph)")
bullet("Draws n = |P| (number of predictions)")
bullet("Observed successes k = |P ∩ GT| = TP")
body("p-value = P(X ≥ k)  where  X ~ Hypergeometric(N, K, n)")
body(
    "All p-values are corrected for multiple testing using the Benjamini–Hochberg "
    "false discovery rate procedure (FDR). A result is reported as significant at "
    "FDR < 0.05."
)

h2("5.3  Fold-enrichment over random")

body(
    "Fold-enrichment expresses how many times more precise the program is compared "
    "to random selection of the same number of genes:"
)
body("Fold-enrichment  =  Precision  /  (|GT| / N)")
body(
    "A fold-enrichment of 10× means the program is 10 times more likely to name a "
    "true target than random chance. This is the most intuitive metric to report "
    "in the results section."
)

h2("5.4  Macro- vs micro-averaged aggregation")

body(
    "When reporting metrics across all 19 test genes:"
)
bullet("Macro-average: compute precision/recall/F1 per gene, then average equally "
       "across genes. Treats all genes equally regardless of how many predictions "
       "they produce. Use this for the main table.")
bullet("Micro-average: pool all TP, FP, and FN across all genes, then compute "
       "metrics once. Dominated by genes with many predictions (e.g. hy5, GRF1). "
       "Include as a secondary table.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — STEP-BY-STEP EXECUTION
# ══════════════════════════════════════════════════════════════════════════════
h1("6.  Step-by-Step Execution")

h2("Step 1  —  Install dependencies  (5 minutes)")
body("The evaluation scripts require scipy and statsmodels in addition to the "
     "existing environment. Run once:")
code("pip install scipy statsmodels pandas")

h2("Step 2  —  Run eval_fw_predictions.py  (several hours)")
body(
    "This script runs the forward BN analysis for all 19 test genes at hops 1–5 "
    "and saves results to Scripts/eval_results/predictions.json. "
    "It saves after each (gene, hop) pair so it can be safely interrupted and "
    "resumed without losing completed work."
)
code("cd bachelorprojekt/Scripts")
code("python eval_fw_predictions.py")
body("Expected runtime:")
bullet("Hop 1–2: seconds to minutes per gene (BoNesis succeeds quickly)")
bullet("Hop 3–4: minutes per gene; BoNesis may time out and fall back to simulation")
bullet("Hop 5: simulation fallback likely for most genes; still produces results")
note(
    "The script records which solver was used at each (gene, hop) in the JSON. "
    "eval_fw_metrics.py can then filter or stratify results by solver if needed."
)

h2("Step 3  —  Run eval_fw_metrics.py  (< 1 minute)")
body("Once predictions.json exists, compute all metrics:")
code("python eval_fw_metrics.py")
body("This produces:")
bullet("eval_results/metrics.csv  —  full table, one row per (gene, hop, prediction type)")
bullet("eval_results/metrics_summary.csv  —  best-hop summary per gene")
body("Summary statistics are also printed directly to the terminal.")

h2("Step 4  —  Open the CSVs in Excel or Python for visualisation")
body("The metrics.csv file is designed for direct import into Excel, R, or a "
     "Jupyter notebook. Key plots to generate:")
bullet("Figure 1: Precision-recall scatter (x = recall, y = precision), one point per "
       "hop per gene, coloured by prediction type. This is the main figure.")
bullet("Figure 2: F1 vs hop depth (line chart, averaged across all genes). Shows "
       "the optimal hop depth.")
bullet("Figure 3: Fold-enrichment bar chart per gene at the optimal hop (shows which "
       "genes are best predicted by the model).")
bullet("Figure 4: Robust vs permissive precision/recall comparison at hops 1–5.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — INDEPENDENT VALIDATION
# ══════════════════════════════════════════════════════════════════════════════
h1("7.  Independent Validation Using RNA-seq Data")

body(
    "The ground-truth comparison above uses the same knowledge graph that the network "
    "was built from — this creates a potential circularity. The most rigorous validation "
    "for a publication uses an independent experimental dataset not used to build the "
    "knowledge graph."
)

h2("7.1  What to look for")
body("Search NCBI GEO (ncbi.nlm.nih.gov/geo) for:")
code("(MYB46 OR NST1 OR NST3) AND Arabidopsis AND RNA-seq")
body(
    "The ideal dataset is an RNA-seq experiment comparing a MYB46 overexpression "
    "line (or NST1/NST3 overexpression, which activates MYB46) with wild-type "
    "Arabidopsis. Genes significantly upregulated in the overexpression condition "
    "are independent evidence for the forward predictions."
)
note(
    "Known published datasets to check: Kim et al. 2013 (Plant Cell, MYB46 "
    "overexpression); McCarthy et al. 2021; Ko et al. 2014. Check that the dataset "
    "is NOT one of the papers that was used to build the knowledge graph — if it is, "
    "it is not independent."
)

h2("7.2  How to use the RNA-seq data")
body("Download the processed differential expression table (DESeq2 or edgeR output). "
     "Define the RNA-seq ground truth as:")
bullet("Activated ground truth: genes with log2FoldChange > 1 and FDR-adjusted p < 0.05")
bullet("Suppressed ground truth: genes with log2FoldChange < −1 and FDR-adjusted p < 0.05")
body(
    "Apply the same clean_name() normalisation and compute precision/recall/F1 "
    "exactly as in Section 5, substituting the RNA-seq gene sets for gt_act / gt_sup."
)

h2("7.3  Why this matters")
body(
    "If the program's perm_activated set overlaps significantly with genes "
    "upregulated in MYB46 overexpression RNA-seq (which was never part of the "
    "training network), you have orthogonal experimental evidence for the method's "
    "biological accuracy. This transforms the work from a methods paper into a "
    "results paper with experimental support."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — WHAT TO REPORT IN THE PAPER
# ══════════════════════════════════════════════════════════════════════════════
h1("8.  What to Report in the Paper")

h2("8.1  Minimum required results")
body("A reviewer will expect to see, for the forward analysis:")
bullet("A table of precision, recall, and F1 at hops 1–5 for the primary gene of "
       "interest (MYB46) and aggregate across all 19 test genes.")
bullet("A precision-recall curve showing the trade-off as hop depth increases.")
bullet("Statistical significance (hypergeometric p-value, FDR-corrected) for each "
       "data point.")
bullet("Comparison to random baseline (fold-enrichment).")
bullet("Comparison of robust_activated vs perm_activated precision and recall.")

h2("8.2  Recommended table structure")

add_table(
    ["Gene", "Hop", "Pred. type", "TP", "Precision", "Recall", "F1",
     "Fold-enr.", "FDR p"],
    [
        ["MYB46", "2", "robust_activated", "—", "—", "—", "—", "—", "—"],
        ["MYB46", "2", "perm_activated",   "—", "—", "—", "—", "—", "—"],
        ["MYB46", "3", "robust_activated", "—", "—", "—", "—", "—", "—"],
        ["MYB46", "3", "perm_activated",   "—", "—", "—", "—", "—", "—"],
        ["All genes (macro-avg)", "3", "perm_activated",
         "—", "—", "—", "—", "—", "—"],
    ],
    col_widths=[3.2, 1.2, 3.8, 1.2, 2.2, 2.0, 1.8, 2.2, 1.8]
)
note("Fill in values from metrics.csv after running the evaluation scripts. "
     "The dashes are placeholders.")

h2("8.3  Key statements to include in the Results section")

body("Based on expected findings (exact numbers from your data):")
bullet(
    "\"At hop 2, the model identifies X% of known MYB46 downstream targets with "
    "Y-fold enrichment over random selection (hypergeometric p = ..., FDR-corrected). "
    "Robust predictions (active in both dark and permissive backgrounds) achieve "
    "Z% precision, compared to W% for permissive-only predictions, confirming that "
    "context-independence is a reliable filter for high-confidence targets.\""
)
bullet(
    "\"Across all 19 test transcription factors, the model achieves a macro-averaged "
    "F1 of X at hop N, significantly outperforming random selection (p < 0.05 in "
    "M/19 genes after FDR correction).\""
)

h2("8.4  Methods section requirements")

body("The following must be explicitly stated in the Methods section:")
bullet("Boolean rule formulation: activators are combined with OR; suppressors with "
       "AND NOT. Any suppressor present blocks activation regardless of the number "
       "of activators. This is a conservative, commonly used formulation "
       "(Thomas 1991; Kauffman 1969).")
bullet("Update scheme: synchronous (all genes updated simultaneously each step). "
       "This is known to produce spurious cyclic attractors not present under "
       "asynchronous update (Fauré et al. 2006, Bioinformatics).")
bullet("Attractor solver: BoNesis (Paulevé et al.) computes all attractors reachable "
       "from a given starting state. When BoNesis exceeds the per-hop timeout, "
       "synchronous simulation is used, which finds exactly one attractor per "
       "starting condition. The solver used for each result is recorded in the output.")
bullet("Ground truth: literature-mined gene–gene regulatory relationships from the "
       "knowledge graph, filtered to pairs where both genes are present in the "
       "network. The ground truth is incomplete; reported recall values are therefore "
       "lower bounds on true recall.")
bullet("Statistical testing: hypergeometric test with Benjamini–Hochberg FDR "
       "correction at a threshold of 0.05.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TIMELINE
# ══════════════════════════════════════════════════════════════════════════════
h1("9.  Realistic Timeline")

add_table(
    ["Day", "Task", "Output"],
    [
        ["1", "Install dependencies; run eval_fw_predictions.py "
              "(leave overnight if needed)",
         "predictions.json"],
        ["2", "Run eval_fw_metrics.py; inspect terminal output; "
              "open metrics.csv in Excel",
         "metrics.csv, metrics_summary.csv"],
        ["2–3", "Generate figures (precision-recall curve, F1 vs hop, "
                "fold-enrichment bar chart) in Excel or Python (matplotlib/seaborn)",
         "3–4 publication figures"],
        ["3", "Search NCBI GEO for independent MYB46 RNA-seq dataset; "
              "download and inspect",
         "RNA-seq ground truth gene list"],
        ["4", "Adapt eval_fw_metrics.py to use RNA-seq ground truth; "
              "re-run and compare",
         "Independent validation metrics"],
        ["4–5", "Write Methods and Results sections using numbers from the CSV tables",
         "Draft manuscript sections"],
    ],
    col_widths=[1.5, 8.5, 5.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — FILE REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("10.  File Reference")

add_table(
    ["File", "Location", "Purpose"],
    [
        ["eval_fw_predictions.py",
         "Scripts/",
         "Data collection: runs BN analysis for 19 genes × 5 hops, saves JSON"],
        ["eval_fw_metrics.py",
         "Scripts/",
         "Metric computation: loads JSON + GT, outputs CSV tables"],
        ["predictions.json",
         "Scripts/eval_results/",
         "Structured prediction sets (auto-created by step 1)"],
        ["metrics.csv",
         "Scripts/eval_results/",
         "Full metrics table, one row per (gene, hop, pred type)"],
        ["metrics_summary.csv",
         "Scripts/eval_results/",
         "Best-hop summary per gene"],
        ["Ground_truth_experiment.csv",
         "raw_data_used_by_scripts/",
         "Source of biological ground truth (207 930 rows)"],
        ["filtered_networkL_normalized.csv",
         "networks_used_by_scripts/",
         "The knowledge-graph network used by the BN program"],
    ],
    col_widths=[4.5, 4.5, 6.5]
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Word document saved to:\n  {OUT_PATH}")
