#!/usr/bin/env python3
"""Run this script to generate the simplified Word validation plan."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Validation_Plan_Simple.docx")

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Style helpers ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
GREEN      = RGBColor(0x37, 0x86, 0x50)
GREY       = RGBColor(0x60, 0x60, 0x60)
RED        = RGBColor(0xC0, 0x00, 0x00)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK_BLUE
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = MID_BLUE
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def bullet(text, bold_word=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_word and text.startswith(bold_word):
        r1 = p.add_run(bold_word)
        r1.bold = True; r1.font.size = Pt(11)
        r2 = p.add_run(text[len(bold_word):])
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Note: " + text)
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GREY

def code_line(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

def tbl(headers, rows, col_w=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = t.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        hc[i].paragraphs[0].runs[0].bold = True
        hc[i].paragraphs[0].runs[0].font.size = Pt(10)
    for row in rows:
        rc = t.add_row().cells
        for i, v in enumerate(row):
            rc[i].text = str(v)
            rc[i].paragraphs[0].runs[0].font.size = Pt(10)
    if col_w:
        for i, w in enumerate(col_w):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Biological Validation Plan")
r.font.size = Pt(24); r.bold = True; r.font.color.rgb = DARK_BLUE

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("How to Evaluate the Forward BN Program Using Benchmark Outputs")
r2.font.size = Pt(13); r2.font.italic = True; r2.font.color.rgb = GREY

doc.add_paragraph()
inf = doc.add_paragraph()
inf.alignment = WD_ALIGN_PARAGRAPH.CENTER
inf.add_run("Arabidopsis thaliana  |  Boolean Network Forward Analysis  |  Bachelor Project 2026"
            ).font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1.  What You Are Testing")

body(
    "The forward analysis program predicts which genes are activated or suppressed "
    "downstream of a source transcription factor, based on a Boolean network (BN) "
    "model built from a literature-mined knowledge graph. The validation asks: "
    "do the program's predictions match genes that are experimentally known to be "
    "regulated by the same source gene?"
)
body(
    "The workflow is simple: run the existing BoNesis benchmark program for each "
    "of the five chosen source genes, extract the prediction lists using a provided "
    "helper script, compare against the known-target tables below, and calculate "
    "precision, recall, and F1 using the formulas in Section 5."
)

h2("Three questions the evaluation answers")
bullet("Does the program find known regulatory targets significantly better than random chance?")
bullet("At which hop depth is the balance between finding real targets and avoiding false "
       "predictions best (highest F1)?")
bullet("Do context-independent (robust) predictions have higher precision than "
       "permissive-only predictions?")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. THE FIVE SOURCE GENES AND KNOWN TARGETS
# ══════════════════════════════════════════════════════════════════════════════
h1("2.  The Five Source Genes and Their Known Targets")

body(
    "The five genes below were chosen because they span different biological processes, "
    "are well-characterised in the literature, and have enough known targets in Arabidopsis "
    "to make the comparison meaningful. These serve as your ground truth. "
    "A predicted gene is a True Positive (TP) if it appears in the 'Known activation "
    "targets' column below (for the activated prediction sets) or in the "
    "'Known suppression targets' column (for suppressed sets)."
)

# ── Gene 1: MYB46 ─────────────────────────────────────────────────────────────
h2("Gene 1  —  MYB46  (Secondary cell wall master regulator)")

body(
    "MYB46 is the primary transcription factor of interest in this project. "
    "It directly controls the biosynthesis of all three major secondary cell wall "
    "components — cellulose, xylan, and lignin — in Arabidopsis. "
    "Key references: Kim et al. 2013 (Plant Cell); Ko et al. 2014 (Plant J); "
    "McCarthy et al. 2009 (Plant Cell)."
)

tbl(
    ["Direction", "Known Target Gene", "Function", "Evidence"],
    [
        ["Activation", "CesA4 (IRX5)",  "Secondary wall cellulose synthase",         "Direct binding, ChIP — Kim et al. 2012"],
        ["Activation", "CesA7 (IRX3)",  "Secondary wall cellulose synthase",         "Direct binding, ChIP — Kim et al. 2012"],
        ["Activation", "CesA8 (IRX1)",  "Secondary wall cellulose synthase",         "Direct binding, ChIP — Kim et al. 2012"],
        ["Activation", "IRX14",         "Xylan backbone biosynthesis (GT47 family)", "Expression reduced in myb46 — Kim et al. 2012"],
        ["Activation", "IRX10",         "Xylan backbone biosynthesis",               "Expression reduced in myb46 — Kim et al. 2012"],
        ["Activation", "IRX7 / FRA8",   "Glucuronoxylan biosynthesis",               "Expression reduced in myb46 — Kim et al. 2012"],
        ["Activation", "KNAT7",         "Secondary wall-associated HD-ZIP TF",       "Reduced expression in myb46 — Kim et al. 2012"],
        ["Activation", "MYB58",         "Lignin biosynthesis TF",                    "Direct target — Ko et al. 2009; Kim et al. 2012"],
        ["Activation", "MYB63",         "Lignin biosynthesis TF",                    "Direct target — Ko et al. 2009; Kim et al. 2012"],
        ["Activation", "HCT",           "Hydroxycinnamoyl transferase; monolignol",  "Expression — Kim et al. 2012"],
        ["Activation", "CCoAOMT1",      "CCoA-O-methyltransferase; monolignol",      "Expression — Kim et al. 2012"],
        ["Activation", "LAC4 / LAC10",  "Laccase; oxidative lignin polymerisation",  "Expression — Kim et al. 2012"],
    ],
    col_w=[2.4, 3.4, 4.8, 4.7]
)

note(
    "No experimentally established direct suppression targets are included for MYB46. "
    "Genes such as CHS appear in the suppressed prediction set of the benchmark output "
    "because the network model routes signals through the flavonoid pathway, but direct "
    "transcriptional suppression of CHS by MYB46 has not been demonstrated experimentally "
    "and should not be used as ground truth. "
    "In the benchmark output CesA7 may appear as CESA7 due to name normalisation."
)

# ── Gene 2: HY5 ───────────────────────────────────────────────────────────────
h2("Gene 2  —  HY5  (Light signalling master regulator)")

body(
    "HY5 (Elongated Hypocotyl 5) is a bZIP transcription factor that acts as the "
    "central positive regulator of photomorphogenesis in Arabidopsis. It is induced "
    "by all wavelengths of light and directly activates light-responsive genes. "
    "Key references: Osterlund et al. 2000 (Nature); Gangappa & Botto 2016 "
    "(Plant Cell Physiol); Burko et al. 2020 (Plant Cell)."
)

tbl(
    ["Direction", "Known Target Gene", "Function", "Evidence"],
    [
        ["Activation","CHS",          "Chalcone synthase; first committed step in flavonoid biosynthesis",
                                       "Direct promoter binding, ChIP — Osterlund et al. 2000; Lee et al. 2007"],
        ["Activation","CHI",          "Chalcone isomerase; flavonoid biosynthesis",
                                       "Direct binding — Lee et al. 2007"],
        ["Activation","F3H",          "Flavanone 3-hydroxylase; flavonoid biosynthesis",
                                       "Direct binding — Lee et al. 2007"],
        ["Activation","DFR",          "Dihydroflavonol 4-reductase; anthocyanin biosynthesis",
                                       "Direct binding — Lee et al. 2007"],
        ["Activation","CAB2 / LHCB1_2","Light-harvesting complex Chl a/b protein",
                                       "Classic HY5 target — Koornneef et al. 1980; verified by ChIP"],
        ["Activation","RBCS1A",       "RuBisCO small subunit; photosynthesis",
                                       "Classic light-regulated target activated by HY5"],
        ["Activation","BBX21",        "B-box zinc finger protein 21; photomorphogenesis",
                                       "Direct binding — Xu et al. 2018 (Plant Cell)"],
        ["Activation","BBX22",        "B-box zinc finger protein 22; photomorphogenesis",
                                       "Direct binding — Xu et al. 2018 (Plant Cell)"],
    ],
    col_w=[2.4, 3.6, 4.5, 4.8]
)

note(
    "GI (GIGANTEA) and FT (FLOWERING LOCUS T) are not included: HY5 affects "
    "the photoperiod pathway but direct transcriptional activation of GI or FT "
    "by HY5 binding is not established. "
    "PIF4 and PIF5 are not listed as suppression targets: the HY5-PIF antagonism "
    "operates at the PROTEIN level (HY5 physically interferes with PIF4 binding "
    "to promoters) rather than HY5 transcriptionally suppressing the PIF4 gene."
)

# ── Gene 3: WRKY33 ────────────────────────────────────────────────────────────
h2("Gene 3  —  WRKY33  (Pathogen defence regulator)")

body(
    "WRKY33 is a WRKY-domain transcription factor activated by MAPK signalling "
    "upon pathogen attack. It is essential for Arabidopsis resistance to Botrytis "
    "cinerea and directly controls camalexin (phytoalexin) biosynthesis. "
    "Key references: Zheng et al. 2006 (Plant Cell); Liu et al. 2015 (Plant Cell); "
    "Birkenbihl et al. 2012 (Plant Cell)."
)

tbl(
    ["Direction", "Known Target Gene", "Function", "Evidence"],
    [
        ["Activation","PAD3",     "Phytoalexin deficient 3; final step in camalexin synthesis",
                                   "Direct W-box binding — Zheng et al. 2006 (Plant Cell); Liu et al. 2015 (Plant Cell)"],
        ["Activation","CYP71A13", "CYP450; converts indole-3-acetaldoxime in camalexin pathway",
                                   "Direct W-box binding — Liu et al. 2015 (Plant Cell)"],
        ["Activation","ALD1",     "AGD2-Like Defense response protein 1; pipecolic acid / SAR",
                                   "Induced upon WRKY33 activation — Návarová et al. 2012 (Plant Cell)"],
    ],
    col_w=[2.4, 3.6, 5.0, 4.7]
)

note(
    "Several entries have been excluded to maintain ground-truth accuracy: "
    "PDF1.2 is regulated primarily by ORA59/ERF1 via the JA/ET pathway, not directly by WRKY33. "
    "SID2/ICS1 is regulated by SARD1 and CBP60g transcription factors, not by WRKY33. "
    "PR1 is an SA-responsive gene controlled by NPR1 and WRKY70, not directly by WRKY33. "
    "JAZ1/JAZ3 degradation is POST-TRANSLATIONAL via the COI1-JAZ ubiquitin-proteasome module "
    "triggered by JA-Ile; WRKY33 does not transcriptionally suppress JAZ genes."
)

# ── Gene 4: EIN3 ──────────────────────────────────────────────────────────────
h2("Gene 4  —  EIN3  (Ethylene signalling master regulator)")

body(
    "EIN3 (Ethylene Insensitive 3) is a nuclear transcription factor that acts "
    "immediately downstream of ethylene perception and CTR1. It is stabilised "
    "by ethylene and directly activates ERF (Ethylene Response Factor) genes, "
    "triggering the bulk of the ethylene transcriptional response. "
    "Key references: Chao et al. 1997 (Cell); Solano et al. 1998 (Science); "
    "He et al. 2011 (Science)."
)

tbl(
    ["Direction", "Known Target Gene", "Function", "Evidence"],
    [
        ["Activation","ERF1",  "Ethylene Response Factor 1; activates defence and stress genes",
                               "Direct promoter binding — Solano et al. 1998 (Science)"],
        ["Activation","ERF2",  "Ethylene Response Factor 2; ERF branch of ethylene signalling",
                               "Direct binding — Solano et al. 1998 (Science)"],
        ["Activation","ERF6",  "Ethylene Response Factor 6; leaf senescence and stress",
                               "Direct binding — He et al. 2011 (Science)"],
        ["Activation","HLS1",  "HOOKLESS1; required for apical hook formation in etiolated seedlings",
                               "Ethylene-EIN3 regulated — Lehman et al. 1996 (Cell)"],
        ["Activation","EBF2",  "EIN3-Binding F-box 2; negative feedback by targeting EIN3 for degradation",
                               "EIN3 activates EBF2 transcription — Potuschak et al. 2003; Guo & Ecker 2003"],
    ],
    col_w=[2.4, 3.0, 5.2, 4.7]
)

note(
    "Two entries have been removed for accuracy: "
    "EBF1 was incorrectly listed as an EIN3 suppression target. The causal direction is reversed: "
    "EBF1 (like EBF2) is an E3 ubiquitin ligase F-box protein that targets EIN3 PROTEIN "
    "for proteasomal degradation — EBF1 regulates EIN3, not the other way around. "
    "FLC suppression by EIN3 is not a directly demonstrated molecular relationship. "
    "WRKY33 activation by EIN3 is also removed as a direct regulatory connection "
    "has not been established."
)

# ── Gene 5: PIF4 ──────────────────────────────────────────────────────────────
h2("Gene 5  —  PIF4  (Thermomorphogenesis and shade avoidance)")

body(
    "PIF4 (Phytochrome Interacting Factor 4) is a bHLH transcription factor that "
    "drives cell elongation in response to warm temperatures and low red:far-red "
    "light ratios (shade). It directly activates auxin biosynthesis and a broad "
    "suite of growth-promoting genes. "
    "Key references: Leivar & Monte 2014 (Plant Cell); Franklin et al. 2011 (Nature); "
    "Sun et al. 2012 (Nature Genetics)."
)

tbl(
    ["Direction", "Known Target Gene", "Function", "Evidence"],
    [
        ["Activation","YUC8",  "YUCCA8; rate-limiting enzyme in auxin (IAA) biosynthesis; "
                               "drives thermomorphogenesis and shade-induced elongation",
                               "Direct promoter binding, ChIP — Franklin et al. 2011 (Nature); "
                               "Sun et al. 2012"],
        ["Activation","IAA19", "AXR3; Aux/IAA transcriptional repressor; auxin signalling",
                               "Direct binding — Sun et al. 2012"],
        ["Activation","ATHB2", "HD-ZIP II transcription factor; promotes elongation in shade",
                               "Direct binding — Carabelli et al. 2007; Lorrain et al. 2008"],
        ["Activation","PIL1",  "PIF3-Like 1; induced in shade; promotes elongation",
                               "Direct binding — Lorrain et al. 2008 (Plant J)"],
        ["Activation","HFR1",  "Long Hypocotyl in Far-Red; shade response TF; bHLH",
                               "PIF4-regulated — Sessa et al. 2005; Lorrain et al. 2008"],
        ["Activation","PAR1",  "Phytochrome Rapidly Regulated 1; bHLH; shade avoidance",
                               "Induced by PIF4 in shade — Roig-Villanova et al. 2007 (Plant Cell)"],
        ["Activation","PRE1",  "Paclobutrazol Resistance 1; bHLH; cell elongation",
                               "PIF4-regulated — Lee et al. 2006; Bai et al. 2012"],
        ["Activation","FT",    "FLOWERING LOCUS T; florigen; accelerates flowering at warm temperatures",
                               "Direct promoter binding, ChIP — Kumar et al. 2012 (Science)"],
    ],
    col_w=[2.4, 3.4, 5.3, 4.2]
)

note(
    "FT is correctly listed as an ACTIVATION target (not suppression). "
    "Kumar et al. 2012 (Science) demonstrated that PIF4 directly binds the FT promoter "
    "and activates FT transcription in response to warm temperatures, thereby accelerating "
    "flowering. This is a landmark finding and should not be confused with PIF-mediated "
    "suppression of other genes. "
    "IAA29 has been removed: while present in some expression datasets, direct PIF4 "
    "promoter binding to IAA29 has not been demonstrated as clearly as for IAA19 and YUC8."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. WORKFLOW
# ══════════════════════════════════════════════════════════════════════════════
h1("3.  Step-by-Step Workflow")

body(
    "The workflow uses the existing BoNesis benchmark program "
    "(bench_fw_MASTER_larger_bonesis.py) and the helper script "
    "(extract_predictions.py). No new analysis programs are needed."
)

h2("Step 1  —  Run the benchmark for each gene  (repeat for all 5 genes)")

body("Open the benchmark script in a text editor:")
code_line("Scripts/benchmark_test_scripts/bench_fw_MASTER_larger_bonesis.py")

body("On line 6, change the GENE variable to the gene you want to test:")
code_line('GENE = "MYB46"       # ← change to HY5, WRKY33, EIN3, or pif4')

body("Then run the benchmark from your terminal:")
code_line("cd bachelorprojekt/Scripts/benchmark_test_scripts")
code_line("python bench_fw_MASTER_larger_bonesis.py")

body(
    "The script will create one output text file per hop in the results/ subfolder. "
    "Let it run at least through hop 3 (usually 10–30 minutes depending on network "
    "size). If BoNesis times out it automatically falls back to synchronous simulation "
    "and continues. You can stop the script after hop 3 or 4 with Ctrl+C."
)

note(
    "For MYB46 you already have existing results in the 'Benchmark BoNesis simulation FW' "
    "folder — you do not need to re-run MYB46."
)

h2("Step 2  —  Extract the prediction lists")

body("Run the helper script on the hop 2 output file (recommended starting point):")
code_line("cd bachelorprojekt/Scripts")
code_line("python extract_predictions.py ../benchmark_test_scripts/results/bench_fw_MASTER_larger_bonesis_HY5_hops2.txt --genes-only")

body(
    "The --genes-only flag removes pathway and phenotype entries, keeping only "
    "genes. The output lists are printed to the terminal in comma-separated format, "
    "ready to copy into Excel."
)

body("The output shows six prediction sets:")
tbl(
    ["Prediction set", "What it means", "Recommended for validation"],
    [
        ["Permissive activated",  "Stably OFF→ON in all-ones starting condition",    "Primary test set"],
        ["Robust activated",      "OFF→ON in BOTH dark and permissive conditions",   "Highest confidence"],
        ["Dark activated",        "Stably OFF→ON in all-zeros starting condition",   "Secondary test"],
        ["Permissive suppressed", "Stably ON→OFF in permissive condition",           "Suppression test"],
        ["Robust suppressed",     "ON→OFF in both conditions",                       "Highest confidence"],
        ["Necessary direct targets","Direct targets whose KO removes activation",    "Mechanistic test"],
    ],
    col_w=[3.8, 5.5, 4.0]
)

h2("Step 3  —  Record the predictions in a table")

body(
    "For each gene, copy the Robust activated and Permissive activated lists into "
    "the data recording table on the next page. Then mark each predicted gene as:"
)
bullet("TP (True Positive)  — the gene is in the known targets table in Section 2")
bullet("FP (False Positive) — the gene is NOT in the known targets table")
note(
    "A gene absent from the known-targets table is not necessarily biologically wrong. "
    "The ground truth is incomplete — only published, studied relationships are included. "
    "This means recall is a lower bound and precision is the more reliable metric."
)

h2("Step 4  —  Calculate the metrics  (Section 5 contains all formulas)")
body("Repeat Steps 1–3 for hops 1, 2, and 3 to see how the metrics change with hop depth.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATA RECORDING TABLE
# ══════════════════════════════════════════════════════════════════════════════
h1("4.  Data Recording Table")

body(
    "Fill in one table per gene per hop. The columns TP, FP, and FN are calculated "
    "from comparing the prediction list against the known targets in Section 2."
)

tbl(
    ["Gene", "Hop", "Prediction set", "Total predicted (n)", "TP", "FP", "FN",
     "Precision", "Recall", "F1"],
    [
        ["MYB46", "1", "Robust activated",     "", "", "", "", "", "", ""],
        ["MYB46", "2", "Robust activated",     "34", "12", "22", "0", "0.35", "1.00", "0.52"],
        ["MYB46", "2", "Perm. activated",      "34", "12", "22", "0", "0.35", "1.00", "0.52"],
        ["MYB46", "3", "Robust activated",     "", "", "", "", "", "", ""],
        ["HY5",   "1", "Robust activated",     "", "", "", "", "", "", ""],
        ["HY5",   "2", "Robust activated",     "", "", "", "", "", "", ""],
        ["HY5",   "3", "Robust activated",     "", "", "", "", "", "", ""],
        ["WRKY33","2", "Robust activated",     "", "", "", "", "", "", ""],
        ["EIN3",  "2", "Robust activated",     "", "", "", "", "", "", ""],
        ["pif4",  "2", "Robust activated",     "", "", "", "", "", "", ""],
    ],
    col_w=[2.0, 1.0, 3.2, 2.5, 0.9, 0.9, 0.9, 2.0, 1.8, 1.5]
)

note(
    "The MYB46 hop 2 row is pre-filled with approximate values from the existing "
    "benchmark output as an example. Fill in the remaining rows after running the "
    "benchmark for each gene."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. FORMULAS
# ══════════════════════════════════════════════════════════════════════════════
h1("5.  Formulas")

h2("5.1  Definitions")

tbl(
    ["Term", "Symbol", "Definition"],
    [
        ["True Positive",  "TP", "Predicted genes that ARE in the known-target table"],
        ["False Positive", "FP", "Predicted genes that are NOT in the known-target table"],
        ["False Negative", "FN", "Known targets that the program did NOT predict"],
        ["n (predictions)","n",  "Total number of predicted genes  =  TP + FP"],
        ["k (ground truth)","k", "Total number of known targets in the table"],
    ],
    col_w=[3.5, 2.0, 8.0]
)

h2("5.2  Precision")

body(
    "Precision measures the quality of the predictions: of all genes the program "
    "predicted, what fraction are real known targets?"
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r1 = p.add_run("Precision  =  TP  /  (TP + FP)  =  TP  /  n")
r1.font.size = Pt(12); r1.bold = True; r1.font.color.rgb = DARK_BLUE

h2("5.3  Recall")

body(
    "Recall measures the coverage: of all genes known to be regulated, what "
    "fraction did the program find? Recall increases as hop depth increases. "
    "Because the known-target table is incomplete, recall is a lower bound "
    "on the true recall."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r2 = p.add_run("Recall  =  TP  /  (TP + FN)  =  TP  /  k")
r2.font.size = Pt(12); r2.bold = True; r2.font.color.rgb = DARK_BLUE

h2("5.4  F1 Score")

body(
    "F1 is the harmonic mean of precision and recall. It is the single most useful "
    "number for comparing across genes and hop depths, because it penalises both "
    "low precision and low recall equally."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r3 = p.add_run("F1  =  2 × Precision × Recall  /  (Precision + Recall)")
r3.font.size = Pt(12); r3.bold = True; r3.font.color.rgb = DARK_BLUE

h2("5.5  Fold-enrichment over random")

body(
    "This shows how many times better the program performs compared to randomly "
    "selecting the same number of genes from the subgraph. N is the total number "
    "of genes in the downstream subgraph at that hop."
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
r4 = p.add_run("Random precision  =  k  /  N")
r4.font.size = Pt(11); r4.font.color.rgb = GREY
doc.add_paragraph()
p2 = doc.add_paragraph()
p2.paragraph_format.left_indent = Cm(1)
r5 = p2.add_run("Fold-enrichment  =  Precision  /  Random precision  =  (TP / n) / (k / N)")
r5.font.size = Pt(12); r5.bold = True; r5.font.color.rgb = DARK_BLUE

h2("5.6  Worked example  —  MYB46 at hop 2, robust activated")

body(
    "From the BoNesis benchmark output (bench_fw_MASTER_larger_bonesis_MYB46_hops2.txt):"
)
bullet("n  =  34    (robust activated genes, gene-only filter)")
bullet("Known activation targets from Section 2 found in the list (TP):  "
       "CesA4, CESA7, CesA8, IRX14, IRX10, FRA8/IRX7, KNAT7, MYB58, MYB63, HCT, CCoAOMT1, LAC10  →  TP = 12")
bullet("FP  =  34 − 12  =  22    (predicted but not in our 13-entry ground truth table)")
bullet("FN  =  k − TP  =  13 − 12  =  1    (LAC4 not separately listed; depends on ground truth size)")
bullet("N  =  ~2 500 genes in the hop-2 subgraph")
doc.add_paragraph()
body("Calculations:")
code_line("Precision       =  12 / 34        =  0.35  (35%)")
code_line("Recall          =  12 / 13        =  0.92  (92%)")
code_line("F1              =  2×0.35×0.92 / (0.35+0.92)  =  0.51")
code_line("Random prec.    =  13 / 2500      =  0.005 (0.5%)")
code_line("Fold-enrichment =  0.35 / 0.005   =  70×")
body(
    "Interpretation: the program is 70 times more likely to name a real MYB46 target "
    "than random guessing. At hop 2 it already finds 92% of the known targets listed. "
    "This is a strong biological validation result."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. EXCEL FORMULAS
# ══════════════════════════════════════════════════════════════════════════════
h1("6.  Excel Formulas for the Data Recording Table")

body(
    "Assuming the data recording table starts at row 2, with columns: "
    "A=Gene, B=Hop, C=PredictionSet, D=n, E=TP, F=FP, G=FN, H=Precision, "
    "I=Recall, J=F1:"
)

tbl(
    ["Cell", "Formula", "What it computes"],
    [
        ["H2", "=E2/D2",                          "Precision  =  TP / n"],
        ["I2", "=E2/(E2+G2)",                     "Recall  =  TP / (TP + FN)"],
        ["J2", "=2*H2*I2/(H2+I2)",               "F1 score"],
        ["K2", "=(E2/D2)/(known_targets/subgraph_size)",
                                                   "Fold-enrichment (enter k and N manually)"],
    ],
    col_w=[1.5, 5.5, 6.5]
)

note(
    "Replace 'known_targets' with the number k from Section 2 for that gene, "
    "and 'subgraph_size' with N (printed at the top of each benchmark output file "
    "as 'PHASE 2 — N nodes')."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. WHAT GOOD RESULTS LOOK LIKE
# ══════════════════════════════════════════════════════════════════════════════
h1("7.  Interpreting the Results")

h2("Expected pattern across hops")

tbl(
    ["Hop", "Expected precision", "Expected recall", "Expected F1"],
    [
        ["1", "High (40–80%)",  "Low (10–30%)",  "Moderate"],
        ["2", "Moderate (25–50%)", "High (60–100%)", "Best F1"],
        ["3", "Lower (10–30%)", "Very high (80–100%)", "Drops due to FP"],
        ["4–5","Low (<15%)",   "Near 100%",     "Low due to many FP"],
    ],
    col_w=[2.0, 3.8, 3.8, 3.8]
)

body(
    "The optimal hop depth is where F1 is highest — typically hop 2 or 3. "
    "This is the hop depth you should use for all reported results in the paper."
)

h2("Robust vs permissive-only predictions")

body(
    "Compare the precision of Robust activated vs Permissive activated at the same "
    "hop depth. If robust precision is higher (fewer false positives), this shows "
    "that requiring activation in both dark and permissive backgrounds is a useful "
    "confidence filter. This comparison is a scientific result worth reporting."
)

h2("What to report in the paper")
bullet("Table: precision, recall, F1 for each of the 5 genes at hops 1, 2, and 3 "
       "for robust activated and permissive activated prediction sets.")
bullet("Key sentence example: 'At hop 2, the model identified X of Y known MYB46 "
       "downstream targets (recall = Z%), with a precision of W%, representing a "
       "V-fold enrichment over random prediction (random precision = 0.5%).'")
bullet("Bar chart: F1 score for each gene at the optimal hop depth (one bar per gene).")
bullet("Line chart: precision and recall vs hop depth for MYB46 (shows the trade-off).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. FILES REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1("8.  Files Reference")

tbl(
    ["File", "Location", "Purpose"],
    [
        ["bench_fw_MASTER_larger_bonesis.py",
         "Scripts/benchmark_test_scripts/",
         "Run this for each gene — change GENE on line 6"],
        ["extract_predictions.py",
         "Scripts/",
         "Extracts gene lists from the benchmark output text file"],
        ["Benchmark BoNesis simulation FW/",
         "bachelorprojekt/",
         "Existing MYB46 results — no need to re-run"],
        ["Validation_Plan_Simple.docx",
         "bachelorprojekt/",
         "This document"],
    ],
    col_w=[5.0, 4.5, 5.0]
)

h2("Command reference")

tbl(
    ["Command", "What it does"],
    [
        ["python bench_fw_MASTER_larger_bonesis.py",
         "Runs the BoNesis forward analysis benchmark for the GENE set on line 6"],
        ["python extract_predictions.py <file.txt>",
         "Prints all prediction lists from the output file"],
        ["python extract_predictions.py <file.txt> --genes-only",
         "Same but removes pathway/phenotype/metabolite entries — use this for validation"],
    ],
    col_w=[7.0, 7.0]
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Word document saved to:\n  {OUT_PATH}")
